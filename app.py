"""
AI Resume & Portfolio Builder
Generate tailored, ATS-aware resumes, cover letters, and portfolio ZIPs from student profile JSON or PDF upload.
Built with Streamlit, classical NLP, and template-based generation.
"""

import streamlit as st
import PyPDF2
import re
import json
import zipfile
from datetime import datetime
from fpdf import FPDF
from collections import Counter
import math

try:
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

try:
    from bs4 import BeautifulSoup
except ImportError:
    BeautifulSoup = None

try:
    import weasyprint
    HAS_WEASYPRINT = True
except ImportError:
    HAS_WEASYPRINT = False


# ======================================================
# PDF and Document Utilities
# ======================================================
def sanitize_html_text(html_content: str) -> str:
    """Extract and sanitize text from HTML"""
    text = re.sub(r'<[^>]+>', '', html_content)
    text = text.replace('&nbsp;', ' ').replace('&amp;', '&').replace('&lt;', '<')
    text = text.replace('&gt;', '>').replace('&quot;', '"')
    text = re.sub(r'\s+', ' ', text)
    return text.strip()



def extract_text_from_pdf(uploaded_file) -> str:
    reader = PyPDF2.PdfReader(uploaded_file)
    text = ""
    for page in reader.pages:
        page_text = page.extract_text()
        if page_text:
            text += page_text + "\n"
    return text.strip()


# ======================================================
# Pure Python TF-IDF Implementation (No scikit-learn)
# ======================================================
def tokenize(text: str) -> list:
    """Tokenize text into words"""
    text = text.lower()
    words = re.findall(r'\b[a-z]+(?:[_-][a-z]+)*\b', text)
    stop_words = {
        'the', 'a', 'an', 'and', 'or', 'but', 'in', 'on', 'at', 'to', 'for',
        'of', 'with', 'by', 'from', 'is', 'are', 'was', 'were', 'be', 'have',
        'has', 'had', 'do', 'does', 'did', 'will', 'would', 'could', 'should',
        'may', 'might', 'must', 'can', 'this', 'that', 'these', 'those', 'i',
        'you', 'he', 'she', 'it', 'we', 'they', 'what', 'which', 'who', 'when',
        'where', 'why', 'how', 'all', 'each', 'every', 'both', 'few', 'more',
        'most', 'other', 'some', 'such', 'no', 'nor', 'not', 'only', 'same',
        'so', 'than', 'too', 'very', 'as', 'if', 'just', 'about', 'into'
    }
    return [w for w in words if w not in stop_words and len(w) > 2]


def extract_ngrams(words: list, n: int = 2) -> list:
    """Extract n-grams from word list"""
    ngrams = []
    for i in range(len(words) - n + 1):
        ngrams.append(' '.join(words[i:i+n]))
    return ngrams


def calculate_tfidf(documents: list) -> dict:
    """Calculate TF-IDF for documents"""
    all_words = []
    doc_tokens = []
    
    for doc in documents:
        tokens = tokenize(doc)
        doc_tokens.append(tokens)
        all_words.extend(tokens)
    
    idf = {}
    total_docs = len(documents)
    
    # Calculate IDF
    for word in set(all_words):
        docs_with_word = sum(1 for tokens in doc_tokens if word in tokens)
        idf[word] = math.log(total_docs / (docs_with_word + 1))
    
    # Calculate TF-IDF
    tfidf_scores = {}
    for i, tokens in enumerate(doc_tokens):
        tf = Counter(tokens)
        for word, count in tf.items():
            tfidf = (count / len(tokens)) * idf.get(word, 0) if tokens else 0
            if tfidf > 0.01:
                if word not in tfidf_scores:
                    tfidf_scores[word] = []
                tfidf_scores[word].append((i, tfidf))
    
    return tfidf_scores


# ======================================================
# ATS Keyword Extraction (Pure Python)
# ======================================================
def extract_keywords(job_description: str) -> list:
    """
    Extract ATS keywords from job description using Pure Python TF-IDF
    
    Args:
        job_description: Job description text
        
    Returns:
        List of extracted keywords
    """
    if not job_description:
        return []
    
    # Common ATS keywords to look for
    common_keywords = [
        'python', 'java', 'javascript', 'sql', 'machine learning', 'data science',
        'cloud', 'aws', 'azure', 'docker', 'kubernetes', 'agile', 'scrum',
        'project management', 'leadership', 'communication', 'analytics',
        'deep learning', 'neural networks', 'nlp', 'computer vision',
        'tensorflow', 'pytorch', 'pandas', 'numpy',
        'git', 'github', 'cicd', 'rest api', 'microservices', 'api', 'database'
    ]
    
    try:
        # Tokenize job description
        tokens = tokenize(job_description)
        
        # Extract 1-grams and 2-grams
        unigrams = tokens
        bigrams = extract_ngrams(tokens, 2)
        all_terms = unigrams + bigrams
        
        # Get TF scores (frequency-based since we only have one doc)
        term_freq = Counter(all_terms)
        sorted_terms = sorted(term_freq.items(), key=lambda x: x[1], reverse=True)
        extracted = [term for term, count in sorted_terms[:20] if count >= 1]
        
        # Also check for common keywords
        jd_lower = job_description.lower()
        found_common = [kw for kw in common_keywords if kw in jd_lower]
        
        # Combine and deduplicate
        all_keywords = list(set(extracted + found_common))
        return all_keywords[:25]
    except Exception as e:
        # Fallback: return common keywords if found
        jd_lower = job_description.lower()
        return [kw for kw in common_keywords if kw in jd_lower][:15]





def merge_keywords(resume_text: str, keywords: list) -> str:
    """Merge ATS keywords naturally into resume text"""
    if not keywords:
        return resume_text
    
    resume_lower = resume_text.lower()
    missing_keywords = [kw for kw in keywords if kw not in resume_lower]
    
    if missing_keywords:
        if 'skills' in resume_lower or 'technical' in resume_lower:
            skills_section = f"\n\nAdditional Skills: {', '.join(missing_keywords[:10])}"
        else:
            skills_section = f"\n\nRelevant Skills: {', '.join(missing_keywords[:10])}"
        return resume_text + skills_section
    
    return resume_text


# ======================================================
# AI / NLP Logic - Enhanced Summaries & STAR Format
# ======================================================
def enhance_resume(resume_text: str, job_description: str = "") -> str:
    """Generate enhanced professional summary using Pure Python"""
    resume_text = " ".join(resume_text.split())

    if not resume_text:
        return (
            "Motivated student seeking opportunities to apply skills, "
            "gain professional experience, and contribute effectively to organizational goals."
        )

    # Extract top keywords using pure Python
    tokens = tokenize(resume_text)
    term_freq = Counter(tokens)
    top_keywords = [word for word, count in term_freq.most_common(8)]

    if not top_keywords:
        top_keywords = ["problem solving", "communication", "teamwork"]

    # Incorporate ATS keywords if job description provided
    if job_description:
        ats_keywords = extract_keywords(job_description)
        relevant_ats = [kw for kw in ats_keywords[:5] if kw not in ' '.join(top_keywords).lower()]
        if relevant_ats:
            top_keywords = top_keywords[:5] + relevant_ats[:3]

    summary = (
        "Results-driven student with practical exposure to "
        + ", ".join(top_keywords[:8])
        + ". Possesses strong analytical thinking, adaptability, and a continuous learning mindset. "
          "Actively seeking opportunities to contribute to impactful projects and grow professionally."
    )

    return summary


def generate_star_bullets(experience_text: str, job_keywords: list = None) -> list:
    """Generate STAR-format bullet points from experience"""
    if not experience_text:
        return []
    
    # Extract action verbs and achievements
    action_verbs = ['developed', 'implemented', 'designed', 'created', 'managed', 'led', 
                   'improved', 'optimized', 'achieved', 'increased', 'reduced', 'delivered']
    
    sentences = re.split(r'[.!?]\s+', experience_text)
    bullets = []
    
    for sentence in sentences[:5]:  # Limit to 5 bullets
        sentence = sentence.strip()
        if len(sentence) > 20 and any(verb in sentence.lower() for verb in action_verbs):
            # Format as STAR bullet
            bullet = f"• {sentence.capitalize()}"
            if not bullet.endswith('.'):
                bullet += '.'
            bullets.append(bullet)
    
    # If no good bullets found, create generic ones
    if not bullets:
        bullets = [
            "• Applied technical skills to solve complex problems and deliver measurable results.",
            "• Collaborated with team members to achieve project objectives and meet deadlines.",
            "• Demonstrated strong problem-solving abilities and attention to detail."
        ]
    
    return bullets[:5]


def generate_cover_letter(name: str, job_description: str, resume_text: str) -> str:
    """Generate a cover letter based on resume and job description"""
    if not job_description:
        return (
            f"Dear Hiring Manager,\n\n"
            f"I am writing to express my interest in the position. "
            f"With my background in {extract_keywords(resume_text)[:3] if resume_text else 'relevant field'}, "
            f"I am confident I would be a valuable addition to your team.\n\n"
            f"Sincerely,\n{name if name else 'Your Name'}"
        )
    
    # Extract key requirements from job description
    keywords = extract_keywords(job_description)
    skills_mentioned = ', '.join(keywords[:5]) if keywords else 'relevant skills'
    
    cover_letter = f"""Dear Hiring Manager,

I am writing to express my strong interest in the position. With my background in {skills_mentioned}, 
I am excited about the opportunity to contribute to your team.

My experience aligns well with the requirements you've outlined. I have a proven track record of 
delivering results and working collaboratively in dynamic environments.

I am eager to discuss how my skills and experience can benefit your organization. Thank you for 
considering my application.

Sincerely,
{name if name else 'Your Name'}"""
    
    return cover_letter


# ======================================================
# PDF Generator with Templates
# ======================================================
def sanitize_text_for_pdf(text: str) -> str:
    """Remove or replace Unicode characters that can't be encoded in latin-1"""
    if not text:
        return ""
    replacements = {
        '\uf10b': '',
        '\u2013': '-',
        '\u2014': '--',
        '\u2018': "'",
        '\u2019': "'",
        '\u201c': '"',
        '\u201d': '"',
        '\u2026': '...',
    }
    result = text
    for unicode_char, replacement in replacements.items():
        result = result.replace(unicode_char, replacement)
    try:
        result.encode('latin-1')
        return result
    except UnicodeEncodeError:
        return result.encode('ascii', 'ignore').decode('ascii')


def create_pdf_resume(name, email, summary, resume_text, template="modern", star_bullets=None) -> str:
    """Create PDF resume with different templates"""
    pdf = FPDF()
    pdf.add_page()
    pdf.set_auto_page_break(auto=True, margin=15)

    safe_name = sanitize_text_for_pdf(name) if name else "Candidate Name"
    safe_email = sanitize_text_for_pdf(email) if email else "email@example.com"
    safe_summary = sanitize_text_for_pdf(summary) if summary else ""
    safe_text = sanitize_text_for_pdf(resume_text.strip() if resume_text else "")

    if template == "modern":
        pdf.set_font("Arial", "B", 18)
        pdf.cell(0, 12, safe_name, ln=True)
        pdf.set_font("Arial", size=11)
        pdf.cell(0, 8, safe_email, ln=True)
        pdf.ln(8)
        pdf.set_font("Arial", "B", 14)
        pdf.cell(0, 10, "Professional Summary", ln=True)
        pdf.set_font("Arial", size=11)
        pdf.multi_cell(0, 6, safe_summary)
        if star_bullets:
            pdf.ln(4)
            pdf.set_font("Arial", "B", 14)
            pdf.cell(0, 10, "Key Achievements", ln=True)
            pdf.set_font("Arial", size=10)
            for bullet in star_bullets[:5]:
                pdf.multi_cell(0, 5, sanitize_text_for_pdf(bullet))
    else:
        pdf.set_font("Times", "B", 16)
        pdf.cell(0, 10, safe_name, ln=True)
        pdf.set_font("Times", size=11)
        pdf.cell(0, 8, safe_email, ln=True)
        pdf.ln(6)
        pdf.set_font("Times", "B", 12)
        pdf.cell(0, 8, "SUMMARY", ln=True)
        pdf.set_font("Times", size=10)
        pdf.multi_cell(0, 6, safe_summary)

    pdf.ln(4)
    pdf.set_font("Arial", "B", 14)
    pdf.cell(0, 10, "Experience & Skills", ln=True)
    pdf.set_font("Arial", size=10)
    if len(safe_text) > 5000:
        safe_text = safe_text[:5000]
    pdf.multi_cell(0, 5, safe_text)

    output_file = f"Resume_{template.capitalize()}.pdf"
    pdf.output(output_file)
    return output_file


# ======================================================
# DOCX Generator
# ======================================================
def create_docx_resume(name, email, summary, resume_text, template="modern", star_bullets=None) -> str:
    """Create DOCX resume"""
    if not HAS_DOCX:
        return None
    
    doc = Document()
    
    safe_name = name if name else "Candidate Name"
    safe_email = email if email else "email@example.com"
    safe_summary = summary if summary else ""
    safe_text = resume_text.strip() if resume_text else ""
    
    heading = doc.add_heading(safe_name, 0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    email_para = doc.add_paragraph(safe_email)
    email_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()
    
    doc.add_heading('Professional Summary', level=1)
    doc.add_paragraph(safe_summary)
    
    if star_bullets:
        doc.add_heading('Key Achievements', level=1)
        for bullet in star_bullets[:5]:
            doc.add_paragraph(bullet, style='List Bullet')
    
    doc.add_heading('Experience & Skills', level=1)
    if len(safe_text) > 5000:
        safe_text = safe_text[:5000]
    doc.add_paragraph(safe_text)
    
    output_file = f"Resume_{template.capitalize()}.docx"
    doc.save(output_file)
    return output_file


# ======================================================
# Portfolio ZIP Generator
# ======================================================
def create_portfolio_zip(name, email, summary, resume_text, cover_letter, pdf_path=None, docx_path=None) -> str:
    """Create a ZIP file containing portfolio documents"""
    zip_filename = f"Portfolio_{name.replace(' ', '_') if name else 'Candidate'}_{datetime.now().strftime('%Y%m%d')}.zip"
    
    with zipfile.ZipFile(zip_filename, 'w', zipfile.ZIP_DEFLATED) as zipf:
        if pdf_path:
            zipf.write(pdf_path, f"Resume_{name.replace(' ', '_') if name else 'Resume'}.pdf")
        if docx_path:
            zipf.write(docx_path, f"Resume_{name.replace(' ', '_') if name else 'Resume'}.docx")
        if cover_letter:
            zipf.writestr("Cover_Letter.txt", cover_letter)
        if summary:
            zipf.writestr("Professional_Summary.txt", summary)
        if resume_text:
            zipf.writestr("Resume_Content.txt", resume_text)
    
    return zip_filename


# ======================================================
# JSON Profile Management
# ======================================================
def save_profile_to_json(name, email, resume_text, job_description="", summary="", cover_letter="") -> str:
    """Save profile data to JSON file"""
    profile = {
        "name": name,
        "email": email,
        "resume_text": resume_text,
        "job_description": job_description,
        "summary": summary,
        "cover_letter": cover_letter,
        "created_at": datetime.now().isoformat()
    }
    
    filename = f"profile_{name.replace(' ', '_') if name else 'candidate'}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.json"
    with open(filename, 'w', encoding='utf-8') as f:
        json.dump(profile, f, indent=2, ensure_ascii=False)
    
    return filename


# ======================================================
# Streamlit UI
# ======================================================
st.set_page_config(
    page_title="AI Resume & Portfolio Builder",
    page_icon="📝",
    layout="wide"
)

st.title("📝 AI Resume & Portfolio Builder")
st.caption("Generate tailored, ATS-aware resumes, cover letters, and portfolio ZIPs")

st.info(
    "This application enhances resumes using **classical NLP techniques** and ATS keyword optimization. "
    "Supports multiple templates and export formats (PDF, DOCX, or both)."
)

# Sidebar options
with st.sidebar:
    st.header("⚙️ Options")
    template_choice = st.selectbox("Choose Template", ["modern", "classic", "professional"])
    export_options = ["PDF"] + (["DOCX", "BOTH"] if HAS_DOCX else [])
    export_format = st.selectbox("Export Format", export_options)
    
    st.header("💾 Profile Management")
    load_profile = st.file_uploader("Load Profile JSON", type=["json"])

# Main content area
tab1, tab2, tab3 = st.tabs(["📄 Upload Resume", "✍️ Create New", "📋 Load Profile"])

profile_data = {}

# Tab 1: Upload Existing Resume
with tab1:
    st.subheader("Upload Existing Resume (PDF)")
    uploaded_file = st.file_uploader("Upload your resume (PDF only)", type=["pdf"], key="upload_pdf")
    name = st.text_input("Your Name", key="name1")
    email = st.text_input("Email Address", key="email1")
    job_description = st.text_area("Job Description (Optional - for ATS optimization)", height=100, key="jd1")
    
    if st.button("Enhance Resume", key="enhance1"):
        if not uploaded_file:
            st.error("Please upload a PDF resume.")
        else:
            with st.spinner("Processing resume..."):
                resume_text = extract_text_from_pdf(uploaded_file)
                
                # ATS keyword extraction
                ats_keywords = extract_keywords(job_description) if job_description else []
                if ats_keywords:
                    resume_text = merge_keywords(resume_text, ats_keywords)
                
                # Generate enhanced summary
                enhanced_summary = enhance_resume(resume_text, job_description)
                
                # Generate STAR bullets
                star_bullets = generate_star_bullets(resume_text, ats_keywords)
                
                # Generate cover letter
                cover_letter = generate_cover_letter(name, job_description, resume_text)
                
                # Store in session state
                profile_data = {
                    "name": name,
                    "email": email,
                    "resume_text": resume_text,
                    "summary": enhanced_summary,
                    "cover_letter": cover_letter,
                    "star_bullets": star_bullets,
                    "job_description": job_description
                }
                st.session_state['profile_data'] = profile_data
                
                st.success("Resume processed successfully!")
                
                # Display results
                col1, col2 = st.columns(2)
                
                with col1:
                    st.subheader("AI-Enhanced Professional Summary")
                    st.write(enhanced_summary)
                    
                    if star_bullets:
                        st.subheader("STAR-Format Achievements")
                        for bullet in star_bullets:
                            st.write(bullet)
                
                with col2:
                    st.subheader("Generated Cover Letter")
                    st.text_area("", cover_letter, height=300, key="cl_display1")
                
                # Generate files
                pdf_path = None
                docx_path = None
                
                if export_format in ["PDF", "BOTH"]:
                    pdf_path = create_pdf_resume(name, email, enhanced_summary, resume_text, template_choice, star_bullets)
                
                if export_format in ["DOCX", "BOTH"] and HAS_DOCX:
                    docx_path = create_docx_resume(name, email, enhanced_summary, resume_text, template_choice, star_bullets)
                
                # Create portfolio ZIP
                portfolio_zip = create_portfolio_zip(name, email, enhanced_summary, resume_text, cover_letter, pdf_path, docx_path)
                
                # Save profile JSON
                profile_json = save_profile_to_json(name, email, resume_text, job_description, enhanced_summary, cover_letter)
                
                # Download buttons
                st.subheader("📥 Downloads")
                col1, col2, col3, col4 = st.columns(4)
                
                if pdf_path:
                    with open(pdf_path, "rb") as f:
                        col1.download_button("⬇️ Download PDF", data=f, file_name=pdf_path, mime="application/pdf")
                
                if docx_path:
                    with open(docx_path, "rb") as f:
                        col2.download_button("⬇️ Download DOCX", data=f, file_name=docx_path, mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
                
                with open(portfolio_zip, "rb") as f:
                    col3.download_button("⬇️ Download Portfolio ZIP", data=f, file_name=portfolio_zip, mime="application/zip")
                
                with open(profile_json, "rb") as f:
                    col4.download_button("⬇️ Download Profile JSON", data=f, file_name=profile_json, mime="application/json")

# Tab 2: Create New Resume
with tab2:
    st.subheader("Create New Resume")
    name2 = st.text_input("Full Name", key="name2")
    email2 = st.text_input("Email Address", key="email2")
    skills = st.text_area("Skills (comma separated)", key="skills2")
    projects = st.text_area("Projects / Experience", height=200, key="projects2")
    job_description2 = st.text_area("Job Description (Optional - for ATS optimization)", height=100, key="jd2")
    
    if st.button("Generate Resume", key="generate2"):
        combined_text = f"{skills}\n{projects}".strip()
        
        if not combined_text:
            st.error("Please enter skills and/or projects.")
        else:
            with st.spinner("Generating resume..."):
                # ATS keyword extraction
                ats_keywords = extract_keywords(job_description2) if job_description2 else []
                
                # Generate enhanced summary
                enhanced_summary = enhance_resume(combined_text, job_description2)
                
                # Generate STAR bullets
                star_bullets = generate_star_bullets(combined_text, ats_keywords)
                
                # Generate cover letter
                cover_letter = generate_cover_letter(name2, job_description2, combined_text)
                
                # Store in session state
                profile_data = {
                    "name": name2,
                    "email": email2,
                    "resume_text": combined_text,
                    "summary": enhanced_summary,
                    "cover_letter": cover_letter,
                    "star_bullets": star_bullets,
                    "job_description": job_description2
                }
                st.session_state['profile_data'] = profile_data
                
                st.success("Resume generated successfully!")
                
                # Display results
                col1, col2 = st.columns(2)
                
                with col1:
                    st.subheader("Generated Professional Summary")
                    st.write(enhanced_summary)
                    
                    if star_bullets:
                        st.subheader("STAR-Format Achievements")
                        for bullet in star_bullets:
                            st.write(bullet)
                
                with col2:
                    st.subheader("Generated Cover Letter")
                    st.text_area("", cover_letter, height=300, key="cl_display2")
                
                # Generate files
                pdf_path = None
                docx_path = None
                
                if export_format in ["PDF", "BOTH"]:
                    pdf_path = create_pdf_resume(name2, email2, enhanced_summary, combined_text, template_choice, star_bullets)
                
                if export_format in ["DOCX", "BOTH"] and HAS_DOCX:
                    docx_path = create_docx_resume(name2, email2, enhanced_summary, combined_text, template_choice, star_bullets)
                
                # Create portfolio ZIP
                portfolio_zip = create_portfolio_zip(name2, email2, enhanced_summary, combined_text, cover_letter, pdf_path, docx_path)
                
                # Save profile JSON
                profile_json = save_profile_to_json(name2, email2, combined_text, job_description2, enhanced_summary, cover_letter)
                
                # Download buttons
                st.subheader("📥 Downloads")
                col1, col2, col3, col4 = st.columns(4)
                
                if pdf_path:
                    with open(pdf_path, "rb") as f:
                        col1.download_button("⬇️ Download PDF", data=f, file_name=pdf_path, mime="application/pdf")
                
                if docx_path:
                    with open(docx_path, "rb") as f:
                        col2.download_button("⬇️ Download DOCX", data=f, file_name=docx_path, mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
                
                with open(portfolio_zip, "rb") as f:
                    col3.download_button("⬇️ Download Portfolio ZIP", data=f, file_name=portfolio_zip, mime="application/zip")
                
                with open(profile_json, "rb") as f:
                    col4.download_button("⬇️ Download Profile JSON", data=f, file_name=profile_json, mime="application/json")

# Tab 3: Load Profile
with tab3:
    st.subheader("Load Saved Profile")
    
    if load_profile:
        try:
            profile_json_data = json.load(load_profile)
            st.success("Profile loaded successfully!")
            
            # Pre-fill form
            name3 = st.text_input("Full Name", value=profile_json_data.get("name", ""), key="name3")
            email3 = st.text_input("Email Address", value=profile_json_data.get("email", ""), key="email3")
            resume_text3 = st.text_area("Resume Text", value=profile_json_data.get("resume_text", ""), height=200, key="resume3")
            job_description3 = st.text_area("Job Description", value=profile_json_data.get("job_description", ""), height=100, key="jd3")
            
            if st.button("Regenerate Resume", key="regenerate3"):
                with st.spinner("Regenerating resume..."):
                    # ATS keyword extraction
                    ats_keywords = extract_keywords(job_description3) if job_description3 else []
                    
                    # Generate enhanced summary
                    enhanced_summary = enhance_resume(resume_text3, job_description3)
                    
                    # Generate STAR bullets
                    star_bullets = generate_star_bullets(resume_text3, ats_keywords)
                    
                    # Generate cover letter
                    cover_letter = generate_cover_letter(name3, job_description3, resume_text3)
                    
                    st.success("Resume regenerated!")
                    
                    # Display and download options (same as above)
                    col1, col2 = st.columns(2)
                    
                    with col1:
                        st.subheader("Professional Summary")
                        st.write(enhanced_summary)
                        
                        if star_bullets:
                            st.subheader("STAR-Format Achievements")
                            for bullet in star_bullets:
                                st.write(bullet)
                    
                    with col2:
                        st.subheader("Cover Letter")
                        st.text_area("", cover_letter, height=300, key="cl_display3")
                    
                    # Generate files
                    pdf_path = None
                    docx_path = None
                    
                    if export_format in ["PDF", "BOTH"]:
                        pdf_path = create_pdf_resume(name3, email3, enhanced_summary, resume_text3, template_choice, star_bullets)
                    
                    if export_format in ["DOCX", "BOTH"] and HAS_DOCX:
                        docx_path = create_docx_resume(name3, email3, enhanced_summary, resume_text3, template_choice, star_bullets)
                    
                    portfolio_zip = create_portfolio_zip(name3, email3, enhanced_summary, resume_text3, cover_letter, pdf_path, docx_path)
                    profile_json = save_profile_to_json(name3, email3, resume_text3, job_description3, enhanced_summary, cover_letter)
                    
                    st.subheader("📥 Downloads")
                    col1, col2, col3, col4 = st.columns(4)
                    
                    if pdf_path:
                        with open(pdf_path, "rb") as f:
                            col1.download_button("⬇️ Download PDF", data=f, file_name=pdf_path, mime="application/pdf")
                    
                    if docx_path:
                        with open(docx_path, "rb") as f:
                            col2.download_button("⬇️ Download DOCX", data=f, file_name=docx_path, mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
                    
                    with open(portfolio_zip, "rb") as f:
                        col3.download_button("⬇️ Download Portfolio ZIP", data=f, file_name=portfolio_zip, mime="application/zip")
                    
                    with open(profile_json, "rb") as f:
                        col4.download_button("⬇️ Download Profile JSON", data=f, file_name=profile_json, mime="application/json")
        
        except json.JSONDecodeError:
            st.error("Invalid JSON file. Please check the file format.")
    else:
        st.info("Upload a JSON profile file to load your saved profile data.")
